#!/usr/bin/env python3
"""
HTML to DOCX Converter

This script converts HTML files to DOCX format with specific formatting requirements:
1. Identifies blocks without text content and skips them
2. Identifies bookmarks and adds them
3. Identifies paragraph types (toc/p/h1/h2/h3/etc) and applies appropriate formatting
4. Links bookmarks during the next written content
5. Adds section page breaks for h2 headings (after TOC is encountered)
6. Creates PageRef form fields for TOC entries
7. Configures page and styles according to specifications
"""

import os
import re
import sys
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION, WD_ORIENTATION
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn


def setup_document():
    """
    Create and configure a new document with specified styles and margins.
    
    Returns:
        Document: Configured docx Document object
    """
    doc = Document()
    
    # Configure sections and margins (mirrored)
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)  # Inside margin when mirrored
    section.right_margin = Inches(0.75)  # Outside margin when mirrored
    section.header_distance = Inches(0.4)
    section.different_first_page_header_footer = False
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    
    # Set up mirrored margins
    section_xml = section._sectPr
    mirror_margins = OxmlElement('w:mirrorMargins')
    section_xml.append(mirror_margins)
    
    # Configure normal style (Book Antiqua 10pt, full justification, 1.15 spacing, 11.5pt after)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Book Antiqua'
    font.size = Pt(10)
    paragraph_format = style.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = 1.15
    paragraph_format.space_after = Pt(11.5)
    
    # Configure heading styles
    for i in range(1, 4):  # h1, h2, h3
        style = doc.styles[f'Heading {i}']
        font = style.font
        font.name = 'Book Antiqua'
        # Set a consistent font color for all headings
        font.color.rgb = RGBColor(0, 0, 0)  # Black color
        
        # Break the link to the theme by setting the style's base properties
        # This ensures Word uses the explicitly defined font rather than the theme font
        style_element = style._element
        rPr = style_element.get_or_add_rPr()
        
        # Set the font directly in the XML to override theme settings
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Book Antiqua')
        rFonts.set(qn('w:hAnsi'), 'Book Antiqua')
        rFonts.set(qn('w:cs'), 'Book Antiqua')
        rFonts.set(qn('w:eastAsia'), 'Book Antiqua')
        rPr.append(rFonts)
        
        # Increase font size differences
        if i == 1:
            font.size = Pt(18)  # Was 14pt
        elif i == 2:
            font.size = Pt(16)  # Was 12pt
        else:
            font.size = Pt(14)  # Was 11pt
        
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Create TOC style
    if 'TOC' not in doc.styles:
        toc_style = doc.styles.add_style('TOC', 1)  # 1 = paragraph style
        font = toc_style.font
        font.name = 'Book Antiqua'
        font.size = Pt(10)
        paragraph_format = toc_style.paragraph_format
        paragraph_format.left_indent = Inches(0.5)
        paragraph_format.space_after = Pt(6)
    
    return doc


def parse_html(html_file):
    """
    Parse HTML file using BeautifulSoup.
    
    Args:
        html_file (str): Path to the HTML file
        
    Returns:
        BeautifulSoup: Parsed HTML
    """
    with open(html_file, 'r', encoding='utf-8') as file:
        html_content = file.read()
    
    # Parse HTML with BeautifulSoup
    soup = BeautifulSoup(html_content, 'html.parser')
    return soup


def has_text_content(element):
    """
    Check if an element has meaningful text content.
    
    Args:
        element: BeautifulSoup element
        
    Returns:
        bool: True if element has text content, False otherwise
    """
    if element.name in ['br', 'hr']:
        return False
    
    # Get text content, stripping whitespace
    text = element.get_text(strip=True)
    
    # Check if there's any text content
    return bool(text)


def collect_bookmarks(soup):
    """
    Find all bookmark elements in the HTML.
    
    Args:
        soup (BeautifulSoup): Parsed HTML
        
    Returns:
        dict: Dictionary of bookmark IDs and their elements
    """
    bookmarks = {}
    
    # Find all elements with id attribute (potential bookmarks)
    for element in soup.find_all(id=True):
        bookmark_id = element['id']
        bookmarks[bookmark_id] = element
        print(f"Collected bookmark from HTML: {bookmark_id} (element: {element.name})")
    
    return bookmarks


def analyze_content(soup):
    """
    Analyze HTML content and structure it for processing.
    
    Args:
        soup (BeautifulSoup): Parsed HTML
        
    Returns:
        list: List of elements to process
    """
    # Get the body content
    body = soup.body
    
    # Find all relevant elements
    elements = []
    
    # Check if body exists
    if body is None:
        # If no body tag, use the entire soup
        body = soup
    
    # Find all elements in the order they appear in the HTML
    all_elements = []
    
    # Recursively collect all elements in order
    def collect_elements(parent):
        for child in parent.children:
            if child.name:  # Skip text nodes
                all_elements.append(child)
                collect_elements(child)
    
    collect_elements(body)
    
    # Filter to keep only the elements we're interested in
    elements = []
    for element in all_elements:
        # Keep headings and paragraphs with text content
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p'] and has_text_content(element):
            elements.append(element)
        # Keep <a> elements with IDs
        elif element.name == 'a' and element.get('id'):
            print(f"Adding <a> element with ID to elements list: {element['id']}")
            elements.append(element)
    
    return elements


# Counter for unique bookmark IDs
bookmark_id_counter = 0

def add_bookmark(paragraph, bookmark_id):
    """
    Add a bookmark to a paragraph.
    
    Args:
        paragraph: docx Paragraph object
        bookmark_id (str): Bookmark ID
        
    Returns:
        None
    """
    global bookmark_id_counter
    
    # Ensure bookmark_id is valid for Word (no spaces or special characters)
    # Replace any invalid characters with underscores
    safe_bookmark_id = re.sub(r'[^\w]', '_', bookmark_id)
    
    # Create a unique ID for this bookmark
    unique_id = str(bookmark_id_counter)
    bookmark_id_counter += 1
    
    # Add the bookmark start
    run = paragraph.add_run()
    tag = run._r
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), unique_id)
    start.set(qn('w:name'), safe_bookmark_id)
    tag.append(start)
    
    # Add the bookmark end
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), unique_id)
    end.set(qn('w:name'), safe_bookmark_id)
    tag.append(end)
    
    print(f"Added bookmark to document: {safe_bookmark_id} (ID: {unique_id})")


def add_page_ref(paragraph, bookmark_id):
    """
    Add a page reference field to a paragraph.
    
    Args:
        paragraph: docx Paragraph object
        bookmark_id (str): Bookmark ID to reference
        
    Returns:
        None
    """
    run = paragraph.add_run()
    
    # Add space before page number
    run.add_text(" ")
    
    # Create the field code for page reference
    r_element = run._r
    
    # Begin field character
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    r_element.append(fldChar1)
    
    # Field instruction text
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' PAGEREF {bookmark_id} \\h '
    r_element.append(instrText)
    
    # Separate field character (important for field to display correctly)
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    r_element.append(fldChar2)
    
    # Add a run for the actual page number that will be displayed
    run2 = OxmlElement('w:r')
    text = OxmlElement('w:t')
    text.text = "0"  # Placeholder, will be replaced by Word with actual page number
    run2.append(text)
    r_element.append(run2)
    
    # End field character
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r_element.append(fldChar3)


def process_toc(element, doc, bookmarks, pending_bookmarks):
    """
    Process a TOC element.
    
    Args:
        element: BeautifulSoup element
        doc: docx Document
        bookmarks (dict): Dictionary of bookmarks
        pending_bookmarks (list): List of pending bookmarks
        
    Returns:
        tuple: (paragraph, True if TOC was processed)
    """
    # Check if this is a TOC entry
    is_toc = False
    if element.name == 'p' and element.get('class'):
        classes = element.get('class')
        if isinstance(classes, list) and 'toc' in classes:
            is_toc = True
        elif isinstance(classes, str) and 'toc' in classes:
            is_toc = True
    
    # Also check for elements that look like TOC entries (have links with internal references)
    if not is_toc and element.name == 'p':
        link = element.find('a', href=lambda href: href and href.startswith('#'))
        if link:
            is_toc = True
    
    if is_toc:
        paragraph = doc.add_paragraph()
        paragraph.style = 'TOC'
        
        # Get the link if it exists
        link = element.find('a')
        if link and link.get('href'):
            href = link.get('href')
            # Extract bookmark ID from href (remove # if present)
            if href.startswith('#'):
                bookmark_id = href[1:]
                
                # Ensure bookmark_id is valid for Word (no spaces or special characters)
                safe_bookmark_id = re.sub(r'[^\w]', '_', bookmark_id)
                
                # Process the link content properly
                process_paragraph_content(link, paragraph)
                
                # Add page reference to the bookmark
                add_page_ref(paragraph, safe_bookmark_id)
                
                print(f"Added TOC entry with reference to bookmark: {safe_bookmark_id} (Original href: {href})")
                
                return paragraph, True
        
        # If no link found, process the content properly
        process_paragraph_content(element, paragraph)
        return paragraph, True
    
    return None, False


def process_heading(element, doc, level, bookmarks, pending_bookmarks, toc_encountered):
    """
    Process a heading element.
    
    Args:
        element: BeautifulSoup element
        doc: docx Document
        level (int): Heading level (1-6)
        bookmarks (dict): Dictionary of bookmarks
        pending_bookmarks (list): List of pending bookmarks
        toc_encountered (bool): Whether a TOC has been encountered
        
    Returns:
        tuple: (paragraph, False)
    """
    # Check if this is a heading of the specified level
    if element.name == f'h{level}':
        # If it's h2 and we've encountered a TOC, add a section break
        if level == 2 and toc_encountered:
            # Add a section break before the heading
            doc.add_section(WD_SECTION.NEW_PAGE)
            
            # Configure the new section with the same settings as the first section
            section = doc.sections[-1]
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1)  # Inside margin when mirrored
            section.right_margin = Inches(0.75)  # Outside margin when mirrored
            section.header_distance = Inches(0.4)
            section.different_first_page_header_footer = False
            section.page_width = Inches(6)
            section.page_height = Inches(9)
            
            # Set up mirrored margins for the new section
            section_xml = section._sectPr
            mirror_margins = OxmlElement('w:mirrorMargins')
            section_xml.append(mirror_margins)
            
            print(f"Added section break for h2 heading: {element.get_text(strip=True)}")
        
        # Add the heading
        paragraph = doc.add_paragraph()
        paragraph.style = f'Heading {level}'
        
        # Check if this element has an ID (bookmark)
        if element.get('id'):
            bookmark_id = element['id']
            # Ensure bookmark_id is valid for Word
            safe_bookmark_id = re.sub(r'[^\w]', '_', bookmark_id)
            print(f"Adding bookmark from heading element: {bookmark_id} -> {safe_bookmark_id}")
            add_bookmark(paragraph, safe_bookmark_id)
        
        # Add any pending bookmarks
        if pending_bookmarks:
            print(f"Adding {len(pending_bookmarks)} pending bookmarks to heading: {element.get_text(strip=True)[:30]}...")
            for bookmark_id in pending_bookmarks:
                # Ensure bookmark_id is valid for Word
                safe_bookmark_id = re.sub(r'[^\w]', '_', bookmark_id)
                print(f"  Adding pending bookmark: {bookmark_id} -> {safe_bookmark_id}")
                add_bookmark(paragraph, safe_bookmark_id)
            pending_bookmarks.clear()
        
        # Process the heading content properly
        process_paragraph_content(element, paragraph)
        
        return paragraph, False
    
    return None, False


def process_paragraph(element, doc, bookmarks, pending_bookmarks):
    """
    Process a normal paragraph element.
    
    Args:
        element: BeautifulSoup element
        doc: docx Document
        bookmarks (dict): Dictionary of bookmarks
        pending_bookmarks (list): List of pending bookmarks
        
    Returns:
        tuple: (paragraph, False)
    """
    if element.name == 'p':
        paragraph = doc.add_paragraph()
        
        # Check if this element has an ID (bookmark)
        if element.get('id'):
            bookmark_id = element['id']
            # Ensure bookmark_id is valid for Word
            safe_bookmark_id = re.sub(r'[^\w]', '_', bookmark_id)
            print(f"Adding bookmark from paragraph element: {bookmark_id} -> {safe_bookmark_id}")
            add_bookmark(paragraph, safe_bookmark_id)
        
        # Add any pending bookmarks
        if pending_bookmarks:
            print(f"Adding {len(pending_bookmarks)} pending bookmarks to paragraph: {element.get_text(strip=True)[:30]}...")
            for bookmark_id in pending_bookmarks:
                # Ensure bookmark_id is valid for Word
                safe_bookmark_id = re.sub(r'[^\w]', '_', bookmark_id)
                print(f"  Adding pending bookmark: {bookmark_id} -> {safe_bookmark_id}")
                add_bookmark(paragraph, safe_bookmark_id)
            pending_bookmarks.clear()
        
        # Process the paragraph content properly to maintain natural text flow
        # This will handle nested elements and preserve proper text formatting
        process_paragraph_content(element, paragraph)
        
        return paragraph, False
    
    return None, False


def process_paragraph_content(element, paragraph):
    """
    Process the content of a paragraph element, handling nested elements properly.
    
    Args:
        element: BeautifulSoup element
        paragraph: docx Paragraph object
        
    Returns:
        None
    """
    # Process all child nodes of the paragraph
    for child in element.children:
        if child.name == 'br':
            # Skip <br> tags - we want natural word wrapping
            continue
        elif child.name is None:  # Text node
            # Add text content, preserving meaningful whitespace but trimming excessive spaces
            text = child.string
            if text:
                # Replace multiple spaces with a single space
                text = ' '.join(text.split())
                if text.strip():  # Only add if there's non-whitespace content
                    paragraph.add_run(text)
        elif child.name in ['b', 'strong']:
            # Bold text
            run = paragraph.add_run(child.get_text(' ', strip=True))
            run.bold = True
        elif child.name in ['i', 'em']:
            # Italic text
            run = paragraph.add_run(child.get_text(' ', strip=True))
            run.italic = True
        elif child.name == 'a':
            # Link text
            run = paragraph.add_run(child.get_text(' ', strip=True))
            run.underline = True
        else:
            # Other elements - get their text content
            text = child.get_text(' ', strip=True)
            if text:
                paragraph.add_run(text)


def process_document(html_file, output_file):
    """
    Process the HTML document and create a DOCX file.
    
    Args:
        html_file (str): Path to the HTML file
        output_file (str): Path to the output DOCX file
        
    Returns:
        None
    """
    # Setup document
    doc = setup_document()
    
    # Parse HTML
    soup = parse_html(html_file)
    
    # Collect bookmarks
    bookmarks = collect_bookmarks(soup)
    
    # Analyze content
    try:
        elements = analyze_content(soup)
    except Exception as e:
        print(f"Error analyzing content: {e}")
        # Fallback to a simpler approach
        elements = []
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p']):
            if has_text_content(element):
                elements.append(element)
    
    # Initialize variables
    toc_encountered = False
    pending_bookmarks = []
    
    # Collect all TOC references to ensure we have bookmarks for them
    toc_references = set()
    for element in soup.find_all('a', href=lambda href: href and href.startswith('#')):
        if element.get('href'):
            href = element.get('href')
            if href.startswith('#'):
                bookmark_id = href[1:]
                toc_references.add(bookmark_id)
                print(f"Found TOC reference in HTML: {bookmark_id} (href: {href})")
    
    # Process each element
    for element in elements:
        # Check if element has an ID (process it even if it has no text content)
        has_id = element.get('id') is not None
        
        # Skip elements without text content, unless they have an ID
        if not has_text_content(element) and not has_id:
            print(f"  Skipping element without text content or ID: {element.name}")
            continue
        
        # Check if element has an ID (we'll handle it as a bookmark)
        if element.get('id'):
            print(f"  Element has ID: {element['id']} (element: {element.name})")
            
            # If it's a heading or paragraph, we'll add the bookmark when processing it
            if element.name not in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p']:
                pending_bookmarks.append(element['id'])
                print(f"  Added to pending bookmarks: {element['id']} (element: {element.name})")
                print(f"  Current pending bookmarks: {pending_bookmarks}")
                continue
            else:
                print(f"  Element with ID will be processed directly: {element['id']} (element: {element.name})")
        
        # Try processing as TOC
        paragraph, is_toc = process_toc(element, doc, bookmarks, pending_bookmarks)
        if paragraph:
            if is_toc:
                toc_encountered = True
            continue
        
        # Try processing as headings
        for level in range(1, 7):  # h1 to h6
            paragraph, _ = process_heading(element, doc, level, bookmarks, pending_bookmarks, toc_encountered)
            if paragraph:
                break
        
        if paragraph:
            continue
        
        # Try processing as normal paragraph
        paragraph, _ = process_paragraph(element, doc, bookmarks, pending_bookmarks)
        if paragraph:
            continue
        
        # If we get here, the element wasn't processed
        # We could add more handlers for other element types if needed
    
    # Check for missing bookmarks and crash if any are found
    missing_bookmarks = toc_references - set(bookmarks.keys())
    if missing_bookmarks:
        error_message = f"Error: Missing bookmarks: {', '.join(missing_bookmarks)}"
        print(error_message)
        raise ValueError(error_message)
    
    # Save the document
    doc.save(output_file)
    
    # Print summary
    print(f"Document saved as {output_file}")
    print(f"Total elements processed: {len(elements)}")
    print(f"Total bookmarks found: {len(bookmarks)}")
    print(f"Total TOC references: {len(toc_references)}")
    if toc_encountered:
        print("Table of Contents was processed")


def main():
    """
    Main function to run the converter.
    """
    if len(sys.argv) < 3:
        print("Usage: python new_html_to_docx_converter.py input.html output.docx")
        sys.exit(1)
    
    html_file = sys.argv[1]
    output_file = sys.argv[2]
    
    if not os.path.exists(html_file):
        print(f"Error: Input file '{html_file}' not found.")
        sys.exit(1)
    
    try:
        process_document(html_file, output_file)
        print(f"Document saved as {output_file}")
    except Exception as e:
        print(f"Error processing document: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()