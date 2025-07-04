#!/usr/bin/env python3
"""
JSON to DOCX Converter

This script converts JSON book files to DOCX format with specific formatting requirements
according to the typesetting specification. It takes a JSON file containing book data
and creates a properly formatted DOCX document with:

1. Title page
2. Table of contents
3. Properly formatted chapters
4. Headers with page numbers
5. Correct page layout and styles

Usage:
    python json_to_docx_converter.py [input_file] [output_file]

    Default input file: irish-boy.json
    Default output file: irish-boy.docx
"""

import os
import sys
import json
import argparse
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_LEADER, WD_BREAK
from docx.enum.section import WD_SECTION, WD_ORIENTATION
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn


# Configuration options
# Set to True to force blank verso (left-hand/even) pages before each chapter begins on an odd page; if the prior chapter ends on an even page, there will be two blank pages before the new chapter
# Set to False to allow chapters to end on an even page immediately precending a new chapter (without forcing blank pages between them)
# Note: Chapters will still always start on odd (right-hand) pages regardless of this setting
FORCE_BLANK_VERSO_PAGES = False


def create_style(doc, style_name, font_name, font_size, alignment=None,
                 space_before=None, space_after=None, first_line_indent=None, **kwargs):
    """
    Create a document style with specified properties.
    
    Args:
        doc: Document object
        style_name: Name of the style to create
        font_name: Font family name
        font_size: Font size in points
        alignment: Paragraph alignment
        space_before: Space before paragraph in points
        space_after: Space after paragraph in points
        first_line_indent: First line indent in inches
        **kwargs: Additional properties (prefix with font_ or para_ to set font or paragraph properties)
    
    Returns:
        The created style
    """
    style = doc.styles.add_style(style_name, 1)  # 1 = paragraph style
    
    # Configure font
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    
    # Set optional font properties
    for prop, value in kwargs.items():
        if prop.startswith('font_'):
            setattr(font, prop[5:], value)
    
    # Configure paragraph format
    paragraph_format = style.paragraph_format
    
    if alignment is not None:
        paragraph_format.alignment = alignment
    if space_before is not None:
        paragraph_format.space_before = Pt(space_before)
    if space_after is not None:
        paragraph_format.space_after = Pt(space_after)
    if first_line_indent is not None:
        paragraph_format.first_line_indent = Inches(first_line_indent)
    
    # Set optional paragraph format properties
    for prop, value in kwargs.items():
        if prop.startswith('para_'):
            setattr(paragraph_format, prop[5:], value)
    
    return style


def setup_document():
    """
    Create and configure a new document with specified styles and margins.
    
    Returns:
        Document: Configured docx Document object
    """
    doc = Document()
    
    # Enable different odd and even headers at the document level
    doc.settings.odd_and_even_pages_header_footer = True
    
    return doc


def configure_styles(doc):
    """
    Configure document styles according to specifications.
    
    Args:
        doc: docx Document object
    """

    # Configure normal style (Book Antiqua 10pt, justified, 1.15 spacing, 11.5pt after)
    normal_style = doc.styles['Normal']
    font = normal_style.font
    font.name = 'Book Antiqua'
    font.size = Pt(10)
    paragraph_format = normal_style.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph_format.line_spacing = 1.15
    paragraph_format.space_after = Pt(11.5)
    paragraph_format.first_line_indent = Inches(0.25)
    
    # Remove all styles except Normal
    styles_to_remove = [s for s in doc.styles if s.name != 'Normal']
    for style in styles_to_remove:
        style.delete()
    
    # 3. Create custom styles
    
    # Create custom BookTitle style (instead of Title)
    create_style(doc, 'BookTitle', 'Georgia', 16,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                font_bold=True)
    
    # Create custom ChapterNumber style (instead of Heading 1)
    create_style(doc, 'ChapterNumber', 'Georgia', 14,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=36, space_after=14,
                font_bold=True)
    
    # Create custom TOCHeading style for Table of Contents heading
    create_style(doc, 'TOCHeading', 'Georgia', 14,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_before=36, space_after=42,
                font_bold=True)
    
    # Create custom ChapterTitle style (instead of Heading 2)
    create_style(doc, 'ChapterTitle', 'Georgia', 12,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_after=42)
    
    # Create BookSubtitle style
    create_style(doc, 'BookSubtitle', 'Georgia', 12,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_after=36)
    
    # Create TOC Entry style
    toc_style = create_style(doc, 'TOC Entry', 'Book Antiqua', 10,
                            space_after=6,
                            first_line_indent=0.25)
    # Add tab stops
    toc_style.paragraph_format.tab_stops.add_tab_stop(Inches(0.81), WD_ALIGN_PARAGRAPH.LEFT)
    toc_style.paragraph_format.tab_stops.add_tab_stop(Inches(4.0), WD_ALIGN_PARAGRAPH.RIGHT, WD_TAB_LEADER.DOTS)
    
    # Create PageHeader style
    header_style = create_style(doc, 'PageHeader', 'Book Antiqua', 9,
                               font_italic=True)
    # Add tab stop for page number
    header_style.paragraph_format.tab_stops.add_tab_stop(Inches(4.25), WD_ALIGN_PARAGRAPH.RIGHT)
    
    # Create DedicationTo style
    create_style(doc, 'DedicationTo', 'Georgia', 12,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_after=20,
                font_bold=True)
    
    # Create DedicationFrom style
    create_style(doc, 'DedicationFrom', 'Georgia', 12,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_after=20,
                font_italic=True)
    
    # Create DedicationCredits style
    create_style(doc, 'DedicationCredits', 'Georgia', 12,
                alignment=WD_ALIGN_PARAGRAPH.CENTER,
                space_after=20)
    
    # Create Quote style for quoted text
    quote_style = create_style(doc, 'Quote', 'Book Antiqua', 10,
                              first_line_indent=0,
                              font_italic=True)
    quote_style.base_style = doc.styles['Normal']  # Base on Normal style
    # Set indents that aren't part of the standard parameters
    paragraph_format = quote_style.paragraph_format
    paragraph_format.left_indent = Inches(0.25)
    paragraph_format.right_indent = Inches(0.25)


def create_title_page(doc, book_data):
    """
    Create the title page with proper formatting.
    
    Args:
        doc: docx Document object
        book_data: Dictionary containing book metadata
    """
    # Configure title page section (first section)
    section = doc.sections[0]
    
    # Setup title page section with no headers
    setup_section_headers(section, "", center_vertical=True, hide_headers=True, reset_numbering=False)
    
    # Add title
    title_paragraph = doc.add_paragraph()
    title_paragraph.style = 'BookTitle'
    title_run = title_paragraph.add_run(book_data['title'].upper())
    
    # Add subtitle if present
    if 'subtitle' in book_data and book_data['subtitle']:
        subtitle_paragraph = doc.add_paragraph()
        subtitle_paragraph.style = 'BookSubtitle'
        subtitle_paragraph.add_run(book_data['subtitle'])
    
    # Add author with "by" prefix
    if 'author' in book_data and book_data['author']:
        author_paragraph = doc.add_paragraph()
        author_paragraph.style = 'BookSubtitle'
        author_paragraph.add_run(f"by {book_data['author']}")
    

def create_dedication_page(doc, dedication_data):
    """
    Create the dedication page with proper formatting.
    
    Args:
        doc: docx Document object
        dedication_data: Dictionary containing dedication information
    """
    # Skip if no dedication data
    if not dedication_data:
        return
    
    # Create section for dedication page (odd page break ensures it starts on a recto page)
    section = doc.add_section(WD_SECTION.ODD_PAGE)
    
    # Setup dedication section with centered vertical alignment and hidden headers
    setup_section_headers(section, "", center_vertical=True, hide_headers=True, reset_numbering=False)
    
    # Add "To" line if present
    if 'to' in dedication_data and dedication_data['to']:
        to_paragraph = doc.add_paragraph()
        to_paragraph.style = 'DedicationTo'
        to_paragraph.add_run(dedication_data['to'])
    
    # Add "From" line if present
    if 'from' in dedication_data and dedication_data['from']:
        from_paragraph = doc.add_paragraph()
        from_paragraph.style = 'DedicationFrom'
        from_paragraph.add_run(dedication_data['from'])
    
    # Add credits if present
    if 'credits' in dedication_data and dedication_data['credits']:
        # Add a blank paragraph with additinal spacing before credits
        blank_paragraph = doc.add_paragraph()
        blank_paragraph.style = 'DedicationFrom'
        
        # Add each credit
        for credit in dedication_data['credits']:
            credit_paragraph = doc.add_paragraph()
            credit_paragraph.style = 'DedicationCredits'
            credit_paragraph.add_run(credit)


def to_title_case(text):
    """
    Convert text to title case following English writing standards.
    
    Args:
        text: Text to convert
        
    Returns:
        str: Text in proper title case
    """
    # First convert the entire string to lowercase
    text = text.lower()
    
    # Words that should not be capitalized (unless they are the first or last word)
    lowercase_words = {
        'a', 'an', 'the',  # Articles
        'and', 'but', 'or', 'nor', 'yet', 'so',  # Coordinating conjunctions
        'as', 'at', 'by', 'for', 'in', 'of', 'on', 'to', 'with',  # Common prepositions
    }
    
    # Split the text into words
    words = text.split()
    
    # Capitalize the first and last words regardless of what they are
    if words:
        words[0] = words[0].capitalize()
        if len(words) > 1:
            words[-1] = words[-1].capitalize()
    
    # Process the rest of the words
    for i in range(1, len(words) - 1):
        if words[i].lower() not in lowercase_words:
            words[i] = words[i].capitalize()
    
    # Join the words back together
    return ' '.join(words)


def roman_numeral(num):
    """
    Convert integer to Roman numeral.
    
    Args:
        num: Integer to convert
        
    Returns:
        str: Roman numeral representation
    """
    val = [
        1000, 900, 500, 400,
        100, 90, 50, 40,
        10, 9, 5, 4,
        1
    ]
    syms = [
        "M", "CM", "D", "CD",
        "C", "XC", "L", "XL",
        "X", "IX", "V", "IV",
        "I"
    ]
    roman_num = ''
    i = 0
    while num > 0:
        for _ in range(num // val[i]):
            roman_num += syms[i]
            num -= val[i]
        i += 1
    return roman_num


# Counter for unique bookmark IDs
bookmark_id_counter = 0

def add_bookmark(paragraph, bookmark_id):
    """
    Add a bookmark to a paragraph.
    
    Args:
        paragraph: docx Paragraph object
        bookmark_id (str): Bookmark ID
    """
    global bookmark_id_counter
    
    # Ensure bookmark_id is valid for Word (no spaces or special characters)
    safe_bookmark_id = ''.join(c if c.isalnum() else '_' for c in bookmark_id)
    
    # Create a unique ID for this bookmark
    unique_id = str(bookmark_id_counter)
    bookmark_id_counter += 1
    
    # Add the bookmark start
    run = paragraph.add_run()
    tag = run._r
    start = OxmlElement('w:bookmarkStart')
    start.set(qn('w:id'), unique_id)
    start.set(qn('w:name'), safe_bookmark_id)
    tag.append(start)
    
    # Add the bookmark end
    end = OxmlElement('w:bookmarkEnd')
    end.set(qn('w:id'), unique_id)
    end.set(qn('w:name'), safe_bookmark_id)
    tag.append(end)


def add_page_ref(paragraph, bookmark_id):
    """
    Add a page reference field to a paragraph.
    
    Args:
        paragraph: docx Paragraph object
        bookmark_id (str): Bookmark ID to reference
    """
    # Ensure bookmark_id is valid for Word
    safe_bookmark_id = ''.join(c if c.isalnum() else '_' for c in bookmark_id)
    
    run = paragraph.add_run()
    
    # Add space before page number
    run.add_text(" ")
    
    # Create the field code for page reference
    r_element = run._r
    
    # Begin field character
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    r_element.append(fldChar1)
    
    # Field instruction text
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = f' PAGEREF {safe_bookmark_id} \\h '
    r_element.append(instrText)
    
    # Separate field character
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    r_element.append(fldChar2)
    
    # Add a run for the actual page number that will be displayed
    run2 = OxmlElement('w:r')
    text = OxmlElement('w:t')
    text.text = "0"  # Placeholder, will be replaced by Word with actual page number
    run2.append(text)
    r_element.append(run2)
    
    # End field character
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r_element.append(fldChar3)


def create_table_of_contents(doc, chapters):
    """
    Create a table of contents with entries for each chapter.
    
    Args:
        doc: docx Document object
        chapters: List of chapter dictionaries
    """
    # Create section for TOC (odd page break ensures TOC starts on a recto page)
    section = doc.add_section(WD_SECTION.ODD_PAGE)
    
    # Setup TOC section with hidden headers
    setup_section_headers(section, "TABLE OF CONTENTS", hide_headers=True, reset_numbering=False)

    # Add TOC heading with dedicated style
    toc_heading = doc.add_paragraph("TABLE OF CONTENTS")
    toc_heading.style = 'TOCHeading'
    
    # Add TOC entries
    for chapter in chapters:
        # Create bookmark ID for the chapter
        chapter_bookmark_id = f"chapter_{chapter['number']}"
        
        # Add TOC entry
        toc_entry = doc.add_paragraph()
        toc_entry.style = 'TOC Entry'
        
        # Add chapter number and title in the correct format
        # Format: "Chapter I<TAB>The New Home<TAB>1"
        chapter_num = roman_numeral(chapter['number'])
        title_text = to_title_case(chapter['title'])
        
        # Add "I." part (Roman numeral with period)
        toc_entry.add_run(f"{chapter_num}.")
        
        # Add tab
        toc_entry.add_run("\t")
        
        # Add title
        toc_entry.add_run(title_text)
        
        # Add tab before page number
        toc_entry.add_run("\t")
        
        # Add page reference
        add_page_ref(toc_entry, chapter_bookmark_id)
    
    # No section breaks added here - they will be added in process_chapters

def add_page_number_field(paragraph):
    """
    Add a page number field to a paragraph.
    
    Args:
        paragraph: docx Paragraph object
    """
    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)
    
    instrText = OxmlElement('w:instrText')
    instrText.text = " PAGE "
    r.append(instrText)
    
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)


def setup_section_headers(section, chapter_title, center_vertical=False, hide_headers=False, reset_numbering=False):
    """
    Configure a section with proper margins, page size, and headers.
    
    Args:
        section: Section object
        chapter_title: Title of the chapter for this section
        center_vertical: Whether to center content vertically (default: False)
        hide_headers: Whether to completely hide headers (default: False)
        reset_numbering: Whether to reset page numbering to 1 and omit even page header (default: False)
    """
    # Configure section margins and page size
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)  # Inside margin when mirrored
    section.right_margin = Inches(0.75)  # Outside margin when mirrored
    section.header_distance = Inches(0.5)
    section.different_first_page_header_footer = True
    section.odd_and_even_pages_header_footer = True
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    
    # Set up mirrored margins
    section_xml = section._sectPr
    mirror_margins = OxmlElement('w:mirrorMargins')
    section_xml.append(mirror_margins)
    
    # Set vertical alignment if requested
    if center_vertical:
        vert_align = OxmlElement('w:vAlign')
        vert_align.set(qn('w:val'), 'center')
        section_xml.append(vert_align)
    else:
        # Ensure top alignment for non-title pages
        vert_align = OxmlElement('w:vAlign')
        vert_align.set(qn('w:val'), 'top')
        section_xml.append(vert_align)
    
    # Handle page numbering
    if reset_numbering:
        # Reset page numbering to start at 1
        pg_num_type = OxmlElement('w:pgNumType')
        pg_num_type.set(qn('w:start'), '1')
        section_xml.append(pg_num_type)
    else:
        # For continuing numbering, actively remove any pgNumType element if it exists
        # This prevents Word from automatically duplicating it from other sections
        pg_num_types = section_xml.findall(qn('w:pgNumType'))
        if pg_num_types:
            for pg_num_type in pg_num_types:
                section_xml.remove(pg_num_type)
    
    # Unlink headers from previous section
    section.header.is_linked_to_previous = False
    section.first_page_header.is_linked_to_previous = False
    section.even_page_header.is_linked_to_previous = False
    
    # Clear all headers
    headers = [section.first_page_header, section.header, section.even_page_header]
    for header in headers:
        for paragraph in header.paragraphs:
            p = paragraph._p
            p.getparent().remove(p)
    
    # If headers should be hidden, return after clearing
    if hide_headers:
        return
    
    # Set up first page header - only page number, no chapter title
    first_para = section.first_page_header.add_paragraph()
    first_para.style = 'PageHeader'
    
    # Check if this is an even or odd section to determine alignment
    # For even sections, align left; for odd sections, align right
    if section.start_type == WD_SECTION.EVEN_PAGE:
        first_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        first_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add only page number on first page
    add_page_number_field(first_para)
    
    # Set up odd (recto) header - right-aligned chapter name with page number
    odd_para = section.header.add_paragraph()
    odd_para.style = 'PageHeader'
    odd_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add chapter name right-aligned on odd (recto) pages
    odd_para.add_run(chapter_title)
    
    # Add tab
    odd_para.add_run("\t")
    
    # Add page number
    add_page_number_field(odd_para)
    
    # Set up even (verso) header - left-aligned chapter name with page number
    even_para = section.even_page_header.add_paragraph()
    even_para.style = 'PageHeader'
    even_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Add page number first on even (verso) pages
    add_page_number_field(even_para)
    
    # Add tab
    even_para.add_run("\t")
    
    # Add chapter name left-aligned
    even_para.add_run(chapter_title)


def process_chapters(doc, chapters):
    """
    Process each chapter with proper formatting and section breaks.
    
    The behavior of blank verso pages before chapters is controlled by the
    FORCE_BLANK_VERSO_PAGES configuration constant.
    
    Args:
        doc: docx Document object
        chapters: List of chapter dictionaries
    """
    for i, chapter in enumerate(chapters):
        # Get the current chapter's title
        chapter_title = to_title_case(chapter['title'])
        
        # Determine if this is the first page (for reset_numbering)
        first_page = (i == 0)
        
        # If configured to force blank verso pages, or if this is the first chapter, add an even page section break
        if FORCE_BLANK_VERSO_PAGES or first_page:
            # Create section breaks before chapter content
            # First add an even page section break (blank verso page)
            section = doc.add_section(WD_SECTION.EVEN_PAGE)
            
            # Setup headers for blank verso page
            setup_section_headers(section, chapter_title,
                                 hide_headers=first_page)
        
        # Add an odd page section break for the chapter content
        section = doc.add_section(WD_SECTION.ODD_PAGE)
        
        # Setup headers for chapter content
        setup_section_headers(section, chapter_title,
                             reset_numbering=first_page)
        
        # Create bookmark ID for the chapter
        chapter_bookmark_id = f"chapter_{chapter['number']}"
        
        # Add chapter number heading
        chapter_num_para = doc.add_paragraph()
        chapter_num_para.style = 'ChapterNumber'
        chapter_num_para.add_run(f"CHAPTER {roman_numeral(chapter['number'])}")
        
        # Add bookmark for TOC reference
        add_bookmark(chapter_num_para, chapter_bookmark_id)
        
        # Add chapter title
        chapter_title_para = doc.add_paragraph()
        chapter_title_para.style = 'ChapterTitle'
        chapter_title_para.add_run(chapter_title)
        
        # Process paragraphs
        for paragraph_item in chapter['paragraphs']:
            # Check if paragraph is a string or an object
            if isinstance(paragraph_item, str):
                # Regular paragraph
                para = doc.add_paragraph(paragraph_item)
                para.style = 'Normal'
            elif isinstance(paragraph_item, dict) and 'type' in paragraph_item and 'content' in paragraph_item:
                # Special paragraph type
                if paragraph_item['type'] == 'quote':
                    # Quote paragraph
                    para = doc.add_paragraph(paragraph_item['content'])
                    para.style = 'Quote'
                else:
                    # Unknown type, default to normal
                    print(f"Warning: Unexpected paragraph type: {paragraph_item['type']}")
                    para = doc.add_paragraph(paragraph_item['content'])
                    para.style = 'Normal'
            else:
                # Fallback for unexpected format
                print(f"Warning: Unexpected paragraph format: {paragraph_item}")
                continue


def process_document(input_file, output_file):
    """
    Process the JSON document and create a DOCX file.
    
    Args:
        input_file: Path to the JSON input file
        output_file: Path to the output DOCX file
    """
    # Load JSON data
    with open(input_file, 'r', encoding='utf-8') as f:
        data = json.load(f)
    
    # Setup document
    doc = setup_document()
    configure_styles(doc)
    
    # Create title page
    create_title_page(doc, data['book'])
    
    # Create dedication page if present
    if 'dedication' in data:
        create_dedication_page(doc, data['dedication'])
    
    # Create table of contents
    create_table_of_contents(doc, data['chapters'])
    
    # Process chapters
    process_chapters(doc, data['chapters'])
    
    # Save the document
    doc.save(output_file)
    print(f"Document saved as {output_file}")


def main():
    """
    Main function to run the converter.
    """
    # Parse command-line arguments
    parser = argparse.ArgumentParser(description='Convert JSON book to DOCX format')
    parser.add_argument('input_file', nargs='?', default='irish-boy.json', 
                        help='Input JSON file (default: irish-boy.json)')
    parser.add_argument('output_file', nargs='?', default='irish-boy.docx',
                        help='Output DOCX file (default: irish-boy.docx)')
    args = parser.parse_args()
    
    # Check if input file exists
    if not os.path.exists(args.input_file):
        print(f"Error: Input file '{args.input_file}' not found.")
        sys.exit(1)
    
    # Process the document
    try:
        process_document(args.input_file, args.output_file)
    except Exception as e:
        print(f"Error processing document: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()